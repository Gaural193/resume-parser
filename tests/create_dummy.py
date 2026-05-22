import os
import docx

def create_dummy_resume(filename, name, email, phone, location):
    doc = docx.Document()
    doc.add_heading(f'{name} Resume', 0)
    doc.add_paragraph(f'Email: {email}')
    doc.add_paragraph(f'Phone: {phone}')
    doc.add_paragraph(f'Location: {location}')
    doc.add_paragraph('Experience: Software Engineer at Tech Corp.')
    
    doc.save(filename)

if __name__ == '__main__':
    os.makedirs('tests/sample_resumes', exist_ok=True)
    create_dummy_resume('tests/sample_resumes/john_doe.docx', 'John Doe', 'john.doe@example.com', '123-456-7890', 'New York')
    create_dummy_resume('tests/sample_resumes/jane_smith.docx', 'Jane Smith', 'jane.smith@example.com', '987-654-3210', 'San Francisco')
    print("Created dummy resumes.")
